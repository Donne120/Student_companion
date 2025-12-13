"""
Document Processor Service
Extracts text from various document formats and chunks for RAG
"""

import io
import re
from typing import List, Optional, Dict, Any
from datetime import datetime
import uuid

from ..utils import logger, DocumentProcessingError, UnsupportedFileTypeError
from ..models import Document, DocumentChunk, DocumentStats, DocumentStatus


class DocumentProcessor:
    """Service for processing documents into searchable chunks"""
    
    # Chunk configuration
    CHUNK_SIZE = 1000  # Target characters per chunk
    CHUNK_OVERLAP = 200  # Overlap between chunks for context
    
    # Supported file types
    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/msword": "doc",
        "text/plain": "txt",
        "text/markdown": "md",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "text/html": "html",
        "application/json": "json",
    }
    
    def __init__(self):
        self.extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "doc": self._extract_docx,  # Same as docx for python-docx
            "txt": self._extract_text,
            "md": self._extract_text,
            "xlsx": self._extract_excel,
            "html": self._extract_html,
            "json": self._extract_json,
        }
    
    async def process_document(
        self,
        content: bytes,
        filename: str,
        content_type: str,
        document_id: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> tuple[str, List[DocumentChunk], DocumentStats]:
        """
        Process a document: extract text and create chunks
        
        Args:
            content: Document bytes
            filename: Original filename
            content_type: MIME type
            document_id: ID of the document record
            metadata: Additional metadata
        
        Returns:
            Tuple of (full_text, chunks, stats)
        """
        import time
        start_time = time.time()
        
        # Determine file type
        file_type = self.SUPPORTED_TYPES.get(content_type)
        if not file_type:
            # Try to infer from filename
            ext = filename.lower().split(".")[-1] if "." in filename else ""
            file_type = ext if ext in self.extractors else None
        
        if not file_type:
            raise UnsupportedFileTypeError(
                content_type,
                list(self.SUPPORTED_TYPES.keys())
            )
        
        logger.info(
            "Processing document",
            document_id=document_id,
            filename=filename,
            file_type=file_type,
            size=len(content)
        )
        
        # Extract text
        extractor = self.extractors.get(file_type)
        if not extractor:
            raise DocumentProcessingError(f"No extractor for file type: {file_type}")
        
        try:
            text, page_count = await extractor(content)
        except Exception as e:
            logger.error("Text extraction failed", error=str(e))
            raise DocumentProcessingError(f"Failed to extract text: {str(e)}", filename)
        
        if not text or len(text.strip()) == 0:
            raise DocumentProcessingError("Document appears to be empty or unreadable", filename)
        
        # Clean text
        text = self._clean_text(text)
        
        # Create chunks
        chunks = self._create_chunks(
            text=text,
            document_id=document_id,
            metadata={
                "filename": filename,
                "category": metadata.get("category", "general") if metadata else "general",
                "title": metadata.get("title", filename) if metadata else filename,
                **(metadata or {})
            }
        )
        
        # Calculate stats
        processing_time = int((time.time() - start_time) * 1000)
        stats = DocumentStats(
            file_size_bytes=len(content),
            page_count=page_count,
            word_count=len(text.split()),
            chunk_count=len(chunks),
            processing_time_ms=processing_time
        )
        
        logger.info(
            "Document processed",
            document_id=document_id,
            chunks=len(chunks),
            words=stats.word_count,
            processing_time_ms=processing_time
        )
        
        return text, chunks, stats
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)
        
        # Remove control characters
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)
        
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        
        # Remove excessive newlines
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        return text.strip()
    
    def _create_chunks(
        self,
        text: str,
        document_id: str,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """Split text into overlapping chunks"""
        chunks = []
        
        # Split by paragraphs first for natural boundaries
        paragraphs = text.split("\n\n")
        
        current_chunk = ""
        current_start = 0
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # If adding this paragraph exceeds chunk size
            if len(current_chunk) + len(para) > self.CHUNK_SIZE and current_chunk:
                # Save current chunk
                chunks.append(DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    start_char=current_start,
                    end_char=current_start + len(current_chunk),
                    metadata={
                        **metadata,
                        "created_at": datetime.utcnow().isoformat()
                    }
                ))
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-self.CHUNK_OVERLAP:] if len(current_chunk) > self.CHUNK_OVERLAP else current_chunk
                current_start = current_start + len(current_chunk) - len(overlap_text)
                current_chunk = overlap_text + "\n\n" + para
                chunk_index += 1
            else:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
        
        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                id=str(uuid.uuid4()),
                document_id=document_id,
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(current_chunk),
                metadata={
                    **metadata,
                    "created_at": datetime.utcnow().isoformat()
                }
            ))
        
        return chunks
    
    # ==================== EXTRACTORS ====================
    
    async def _extract_pdf(self, content: bytes) -> tuple[str, Optional[int]]:
        """Extract text from PDF"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return "\n\n".join(text_parts), len(reader.pages)
            
        except ImportError:
            # Fallback to PyPDF2
            from PyPDF2 import PdfReader as PdfReader2
            
            reader = PdfReader2(io.BytesIO(content))
            text_parts = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return "\n\n".join(text_parts), len(reader.pages)
    
    async def _extract_docx(self, content: bytes) -> tuple[str, Optional[int]]:
        """Extract text from Word document"""
        from docx import Document as DocxDocument
        
        doc = DocxDocument(io.BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts), None
    
    async def _extract_text(self, content: bytes) -> tuple[str, Optional[int]]:
        """Extract text from plain text file"""
        # Try different encodings
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                text = content.decode(encoding)
                return text, None
            except UnicodeDecodeError:
                continue
        
        raise DocumentProcessingError("Unable to decode text file")
    
    async def _extract_excel(self, content: bytes) -> tuple[str, Optional[int]]:
        """Extract text from Excel file"""
        from openpyxl import load_workbook
        
        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        text_parts = []
        
        for sheet in wb.worksheets:
            sheet_text = [f"Sheet: {sheet.title}"]
            
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(row_values):
                    sheet_text.append(" | ".join(row_values))
            
            text_parts.append("\n".join(sheet_text))
        
        return "\n\n".join(text_parts), len(wb.worksheets)
    
    async def _extract_html(self, content: bytes) -> tuple[str, Optional[int]]:
        """Extract text from HTML"""
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(content, "html.parser")
        
        # Remove script and style elements
        for element in soup(["script", "style", "nav", "footer", "header"]):
            element.decompose()
        
        text = soup.get_text(separator="\n")
        return text, None
    
    async def _extract_json(self, content: bytes) -> tuple[str, Optional[int]]:
        """Extract text from JSON"""
        import json
        
        data = json.loads(content.decode("utf-8"))
        
        def extract_strings(obj, depth=0) -> List[str]:
            strings = []
            if isinstance(obj, dict):
                for key, value in obj.items():
                    strings.append(f"{'  ' * depth}{key}:")
                    strings.extend(extract_strings(value, depth + 1))
            elif isinstance(obj, list):
                for item in obj:
                    strings.extend(extract_strings(item, depth))
            elif isinstance(obj, str):
                strings.append(f"{'  ' * depth}{obj}")
            elif obj is not None:
                strings.append(f"{'  ' * depth}{str(obj)}")
            return strings
        
        text = "\n".join(extract_strings(data))
        return text, None


# Singleton instance
document_processor = DocumentProcessor()

